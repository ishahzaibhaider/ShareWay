"""Generate Assignment 4 docx files for ShareWay group G15."""
from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

MEMBERS = [
    {
        "reg": "FA23-BAI-050",
        "name": "Syed Muhammad Shahzaib Haider Rizvi",
        "module": "Authentication, Profile & Vehicle Management Module",
        "use_case_id": "UC-A.1",
        "use_case_name": "Sign Up / User Registration",
    },
    {
        "reg": "FA23-BAI-006",
        "name": "Arhum Fareed",
        "module": "Ride Management Module",
        "use_case_id": "UC-3.1",
        "use_case_name": "Offer a Ride",
    },
    {
        "reg": "FA23-BAI-028",
        "name": "Muhammad Awais Ali",
        "module": "Communication, History & Reports Module",
        "use_case_id": "UC-4.2",
        "use_case_name": "Submit Rating & Review",
    },
]

GROUP_NO = "15"
PROJECT_TITLE = "ShareWay - AI Powered Peer-to-Peer Ride Sharing Platform"
COURSE = "CSC291 Software Engineering"
SECTION = "BS(AI) - Section 6A"
SEMESTER = "Spring 2026"
SUBMISSION_DATE = "12 - May - 2026"

NAVY = RGBColor(0x1B, 0x29, 0x3F)
MINT = RGBColor(0x2E, 0xC4, 0xA0)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREY = RGBColor(0x55, 0x5B, 0x6E)

PROJECT_OVERVIEW = (
    "ShareWay is a cross-platform mobile application built with Flutter that facilitates "
    "peer-to-peer ride-sharing among urban commuters in Pakistan. The system uses Firebase "
    "for real-time backend services (Authentication, Cloud Firestore, Cloud Messaging) and "
    "OpenStreetMap for all map operations (Nominatim for geocoding, OSRM for routing). A "
    "weighted AI-driven matching algorithm scores potential rides on route similarity "
    "(40 percent), time overlap (30 percent), driver rating (20 percent) and pickup proximity "
    "(10 percent), so that passengers are connected to the most suitable drivers. The "
    "application supports dual user roles (driver and passenger), real-time messaging, live "
    "GPS tracking, ratings & reviews, ride history, and four Pakistani payment methods "
    "(Cash, Easypaisa, JazzCash, Bank Transfer)."
)


# ---------- low level OOXML helpers ----------

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, left=120, bottom=80, right=120):
    """Cell padding in twips (1/20 of a point). 120 twips ~= 6pt."""
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tc_pr.append(mar)


def set_table_borders(table, color="1B293F", size=6):
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size))
        b.set(qn("w:color"), color)
        borders.append(b)
    tbl_pr.append(borders)


def configure_default_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.25

    for level, size in [(1, 18), (2, 14), (3, 12)]:
        style = doc.styles[f"Heading {level}"]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = NAVY
        pf = style.paragraph_format
        pf.space_before = Pt(20 if level == 1 else 14)
        pf.space_after = Pt(10 if level == 1 else 8)
        pf.keep_with_next = True

    sections = doc.sections
    for s in sections:
        s.top_margin = Cm(2.2)
        s.bottom_margin = Cm(2.2)
        s.left_margin = Cm(2.0)
        s.right_margin = Cm(2.0)


# ---------- friendly content builders ----------

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
    return h


def add_para(doc, text, bold=False, italic=False, size=11, align=None, space_after=None, color=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.bold = bool(bold)
    r.italic = bool(italic)
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def add_spacer(doc, pts=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pts)
    return p


def add_table_caption(doc, caption):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(caption)
    r.bold = True
    r.italic = True
    r.font.size = Pt(10)
    r.font.color.rgb = NAVY
    return p


def build_table(doc, headers, rows, header_size=10, body_size=9, col_widths_cm=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(header_size)
        r.font.color.rgb = WHITE
        set_cell_shading(hdr[i], "1B293F")
        set_cell_margins(hdr[i], top=100, bottom=100, left=140, right=140)
        hdr[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = ""
            text_blocks = str(val).split("\n")
            for bi, blk in enumerate(text_blocks):
                if bi == 0:
                    p = cell.paragraphs[0]
                else:
                    p = cell.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.15
                run = p.add_run(blk)
                run.font.size = Pt(body_size)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "F4F6F9")

    if col_widths_cm:
        for row in table.rows:
            for idx, w in enumerate(col_widths_cm):
                row.cells[idx].width = Cm(w)

    set_table_borders(table)
    add_spacer(doc, pts=10)
    return table


def horizontal_rule(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2EC4A0")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ---------- cover page ----------

def cover_page(doc, member):
    add_spacer(doc, pts=20)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMSATS University Islamabad")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Department of Computer Science")
    r.font.size = Pt(14)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(20)

    horizontal_rule(doc)

    for text, size, bold, color, after in [
        (COURSE, 16, True, NAVY, 4),
        (f"{SECTION}   |   Semester: {SEMESTER}", 12, False, GREY, 24),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.color.rgb = color
        p.paragraph_format.space_after = Pt(after)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Assignment No. 4")
    r.bold = True
    r.font.size = Pt(28)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(4)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Use-Case Based Testing, Equivalence Partitioning")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("and Boundary Value Analysis")
    r.italic = True
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(30)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Group No. {GROUP_NO}")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(20)

    horizontal_rule(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Project Title")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(PROJECT_TITLE)
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Submitted By")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(member["name"])
    r.bold = True
    r.font.size = Pt(15)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(member["reg"])
    r.font.size = Pt(13)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(16)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Module:  {member['module']}")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = GREY
    p.paragraph_format.space_after = Pt(2)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Selected Core Process:  {member['use_case_name']}")
    r.bold = True
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = NAVY
    p.paragraph_format.space_after = Pt(28)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Submission Date:  {SUBMISSION_DATE}")
    r.font.size = Pt(11)
    r.font.color.rgb = GREY

    doc.add_page_break()


# ---------- main sections ----------

def section_project_overview(doc):
    add_heading(doc, "1.  Project Overview", level=1)
    add_para(doc, PROJECT_OVERVIEW)


def section_group_split(doc, current_reg):
    add_heading(doc, "2.  Assignment Scope and Group Division", level=1)
    add_para(doc,
             "Assignment 4 requires every member of the group to select a different core "
             "process and exercise it through use-case-based testing, equivalence "
             "partitioning and boundary value analysis. The table below records the "
             "non-overlapping use case picks for the three members of Group 15.")
    add_table_caption(doc, "Table 2.1  -  Use-case allocation across Group 15 members")
    headers = ["Member", "Registration", "Module", "Selected Use Case"]
    rows = [
        [m["name"], m["reg"], m["module"], f"{m['use_case_id']} - {m['use_case_name']}"]
        for m in MEMBERS
    ]
    build_table(doc, headers, rows, col_widths_cm=[4.5, 3.0, 4.5, 4.5])

    current = next(m for m in MEMBERS if m["reg"] == current_reg)
    add_para(
        doc,
        f"This document is the personal contribution of {current['name']} ({current_reg}), "
        f"covering use case {current['use_case_id']}  -  {current['use_case_name']}. "
        "The remainder of the document is organised into four sections: a detailed "
        "specification of the selected use case, the derived use-case-based test "
        "cases, the equivalence partitioning of each input field, and finally the "
        "boundary value analysis.",
        italic=True,
    )


# ---------- per-member content ----------

SIGNUP_DESC = {
    "id": "UC-A.1",
    "name": "Sign Up / User Registration",
    "actors": "New User (Primary), Firebase Authentication (Secondary), Firebase Phone Auth (Secondary)",
    "preconditions": [
        "The user has the ShareWay application installed on a supported Android or iOS device.",
        "The device has an active internet connection.",
        "The user does not already have a registered account with the entered email.",
    ],
    "postconditions": [
        "A new Firebase Authentication record and matching users/{uid} Firestore document are created.",
        "The user's phone number has been verified through a one-time password (OTP).",
        "The user is signed in and navigated to the Find Ride home screen with a welcome banner.",
    ],
    "normal_flow": [
        "User opens the Sign Up screen from the Login screen.",
        "User enters Full Name, Email, Phone Number, Password and Confirm Password.",
        "User selects an Account Type (Passenger | Driver | Both).",
        "User taps Send OTP.",
        "System sends a 6-digit OTP via Firebase Phone Authentication.",
        "User enters the OTP within the 60-second validity window.",
        "User taps Create Account.",
        "System creates the Firebase Auth account, writes the users/{uid} document and stores the FCM token.",
        "System navigates the user to the Find Ride home screen with a welcome banner.",
    ],
    "alt": [
        ("2a Email already in use", [
            "In step 2, if the entered email is already registered, the system displays an inline error 'Email already in use'.",
            "Use case resumes at step 2.",
        ]),
        ("4a Phone number invalid", [
            "In step 4, the Send OTP button remains disabled until the phone field passes +92 format validation.",
            "User corrects the phone and use case resumes at step 4.",
        ]),
    ],
    "exceptions": [
        ("6a OTP incorrect", [
            "In step 6, if the entered OTP does not match, the system displays 'Code incorrect, try again' with a cooldown.",
            "Use case resumes at step 6.",
        ]),
        ("6b OTP expired", [
            "If the user enters the OTP more than 60 seconds after it was issued, the system rejects the code and offers Resend OTP.",
            "Use case resumes at step 4.",
        ]),
        ("8a Network failure", [
            "In step 8, if the network drops, the system shows a snackbar 'No internet connection' and retains the entered data so the user can retry.",
        ]),
    ],
    "quality": "All credentials are managed by Firebase Authentication; the password is never stored by the application. Sign Up should complete within 3-4 seconds on a typical 4G connection after the OTP is verified.",
}

SIGNUP_TCS = [
    ["TC-SU-01", "Successful sign up (normal flow)",
     "App installed; device online; email unused",
     "1. Open Sign Up.\n2. Enter all valid fields.\n3. Tap Send OTP.\n4. Enter correct OTP.\n5. Tap Create Account.",
     "Name='Shahzaib H.'\nEmail='new@test.com'\nPhone='+923001234567'\nPassword='Pass1234'\nType='Passenger'\nOTP=correct",
     "Account created; users/{uid} document written; user lands on Find Ride home with welcome banner.",
     "High"],
    ["TC-SU-02", "Duplicate email rejected",
     "Email 'existing@test.com' already registered",
     "1. Enter the registered email.\n2. Complete other fields.\n3. Tap Create Account.",
     "Email='existing@test.com'",
     "Inline error 'Email already in use'; account not created.",
     "High"],
    ["TC-SU-03", "Weak password (no digit)",
     "All other fields valid",
     "1. Enter password 'password'.\n2. Tap Create Account.",
     "Password='password'",
     "Inline error 'Password must be at least 8 characters and contain a letter and a number'.",
     "High"],
    ["TC-SU-04", "Confirm Password mismatch",
     "All other fields valid",
     "1. Enter Password.\n2. Enter a different Confirm Password.\n3. Tap Create Account.",
     "Password='Pass1234'\nConfirm='Pass4321'",
     "Inline error 'Passwords do not match'; Create Account disabled.",
     "High"],
    ["TC-SU-05", "Invalid email format",
     "All other fields valid",
     "1. Enter malformed email.\n2. Try to submit.",
     "Email='user@@test'",
     "Inline error 'Invalid email format'; Create Account disabled.",
     "Medium"],
    ["TC-SU-06", "OTP incorrect",
     "OTP has been sent within the last 60s",
     "1. Enter a wrong OTP.\n2. Tap Create Account.",
     "OTP='000000' (actual differs)",
     "Error 'Code incorrect, try again' with cooldown; account not created.",
     "High"],
    ["TC-SU-07", "OTP expired (>60s)",
     "OTP issued more than 60 seconds ago",
     "1. Wait 61 seconds.\n2. Enter the original OTP.\n3. Tap Create Account.",
     "OTP age = 61s",
     "Error 'OTP expired'; Resend OTP option shown.",
     "Medium"],
    ["TC-SU-08", "Network dropped during account creation",
     "OTP verified successfully",
     "1. Turn off Wi-Fi and mobile data.\n2. Tap Create Account.",
     "Network off",
     "Snackbar 'No internet connection'; entered data retained for retry.",
     "Medium"],
    ["TC-SU-09", "Full Name below minimum length",
     "All other fields valid",
     "1. Enter a 1-character name.\n2. Tap Create Account.",
     "Name='A'",
     "Inline error 'Name must be 2-50 characters'.",
     "Low"],
    ["TC-SU-10", "Full Name above maximum length",
     "All other fields valid",
     "1. Enter a 51-character name.\n2. Tap Create Account.",
     "Name = 51 'A' characters",
     "Inline error 'Name must be 2-50 characters'.",
     "Low"],
    ["TC-SU-11", "Phone missing country code",
     "All other fields valid",
     "1. Enter 10 digits without the +92 prefix.",
     "Phone='3001234567'",
     "Field auto-prepends +92 or shows inline error 'Phone must start with +92'.",
     "Medium"],
    ["TC-SU-12", "Non-numeric OTP entry",
     "OTP has been sent",
     "1. Try to type letters into the OTP input.",
     "OTP='ABCDEF'",
     "Field rejects non-numeric characters; only digits accepted.",
     "Low"],
    ["TC-SU-13", "Default Account Type is Passenger",
     "Fresh Sign Up screen",
     "1. Open Sign Up.\n2. Observe the segmented control.",
     "(no selection)",
     "Account Type defaults to Passenger.",
     "Low"],
    ["TC-SU-14", "Account Type 'Both' is persisted",
     "User selects Both",
     "1. Select 'Both'.\n2. Complete sign up.",
     "Type='Both'",
     "users/{uid}.role == 'both' in Firestore.",
     "Low"],
]

SIGNUP_EP = [
    ["Full Name (length, chars)", "Length 2 to 50", "Length < 2; Length > 50; empty"],
    ["Email", "Matches RFC-5322 pattern AND not already registered", "Pattern mismatch (missing @, double @, no domain); duplicate email; empty"],
    ["Phone (digits after +92)", "Exactly 10 numeric digits", "< 10 digits; > 10 digits; non-numeric characters; missing +92 prefix; empty"],
    ["Password", "Length >= 8 AND contains at least one letter AND at least one digit", "Length < 8; only letters; only digits; empty"],
    ["Confirm Password", "Exactly equal to Password", "Different from Password; empty"],
    ["OTP Code (digits)", "Exactly 6 numeric digits", "< 6 digits; > 6 digits; non-numeric; empty"],
    ["OTP age (seconds since issue)", "0 to 60", "> 60 (expired)"],
    ["Account Type", "One of {Passenger, Driver, Both}", "Any other value (rejected at API level)"],
]

SIGNUP_BVA = [
    ["Full Name length", "1", "2", "26", "50", "51"],
    ["Password length", "7", "8", "10", "128", "(no enforced upper - 128 used as representative upper)"],
    ["Phone digits (after +92)", "9", "10", "10", "10", "11"],
    ["OTP digits", "5", "6", "6", "6", "7"],
    ["OTP age (seconds)", "n/a", "0", "30", "60 (boundary)", "61"],
]

OFFER_DESC = {
    "id": "UC-3.1",
    "name": "Offer a Ride",
    "actors": "Driver (Primary), OpenStreetMap/OSRM (Secondary), Firebase Firestore (Secondary)",
    "preconditions": [
        "User is authenticated and logged into the application.",
        "User has the Driver (or Both) role active on their profile.",
        "User has at least one vehicle saved in users/{uid}/vehicles.",
        "Location and Internet permissions are granted on the device.",
    ],
    "postconditions": [
        "A new rides document is written to the Firestore rides collection with status = active.",
        "The ride is immediately discoverable by passengers via the Search & Match Ride use case.",
        "For a recurring ride, child rides are spawned for each selected day by a Cloud Function.",
    ],
    "normal_flow": [
        "Driver navigates to the Offer a Ride screen.",
        "Driver selects the origin and destination on the interactive map using the Map Location Picker.",
        "System reverse-geocodes both points (Nominatim) and validates that they fall inside Pakistan.",
        "System calculates the route polyline, distance and duration using OSRM.",
        "Driver inputs departure date, departure time, available seats (1-4) and accepted payment methods.",
        "Driver selects ride type (One-time | Recurring). If Recurring, the driver also selects at least one day of the week.",
        "System calculates the estimated fare as distanceKm * 25 PKR.",
        "Driver taps Create Ride.",
        "System writes the rides document and shows a green snackbar 'Ride created successfully!'.",
    ],
    "alt": [
        ("4a OSRM route calculation fails", [
            "In step 4, if OSRM cannot find a route between the two points, the system prompts the Driver to adjust the map pins.",
            "Driver re-pins origin / destination; use case resumes at step 3.",
        ]),
        ("6a Driver chose recurring ride", [
            "In step 6, the driver toggles Ride Type to Recurring and selects one or more days of the week.",
            "On Create, rideType is stored as recurring and recurringDays is populated.",
            "Use case resumes at step 7.",
        ]),
    ],
    "exceptions": [
        ("2a Geographic constraint violation", [
            "In step 2, if either the origin or the destination falls outside Pakistan, the system displays 'Location must be within Pakistan.'.",
            "Use case resumes at step 2.",
        ]),
        ("5a Departure time too soon", [
            "If the selected departure time is less than 5 minutes from the current time, the system disables Create Ride and shows an inline error 'Departure must be at least 5 minutes in the future.'.",
        ]),
        ("8a Network failure during write", [
            "In step 8, if the device is offline, Firestore offline persistence caches the write locally and replays it when the connection is restored.",
        ]),
    ],
    "quality": "Map interactions and route calculation must remain responsive on a mid-range device. The OSRM round-trip should complete within 3 seconds on a typical 4G connection. Offline-cached writes must flush automatically within 30 seconds of reconnection.",
}

OFFER_TCS = [
    ["TC-OR-01", "Normal flow - publish a valid one-time ride",
     "Driver role active; one vehicle on profile; online",
     "1. Open Offer a Ride.\n2. Pick origin and destination inside Pakistan.\n3. Pick date today+1 and time now+30 min.\n4. Set seats=2; payment={Cash}.\n5. Tap Create Ride.",
     "Origin=Lahore Gulberg\nDestination=DHA Phase 5\nDate=today+1\nTime=now+30 min\nSeats=2\nMethods={Cash}",
     "rides document created with status=active; green snackbar 'Ride created successfully!'.",
     "High"],
    ["TC-OR-02", "Origin outside Pakistan rejected",
     "Driver role active",
     "1. Open Offer a Ride.\n2. In Map Picker, drag origin pin to Dubai.",
     "Origin lat,lng inside UAE",
     "Inline error 'Location must be within Pakistan.'; field not accepted.",
     "High"],
    ["TC-OR-03", "OSRM cannot resolve a route",
     "Origin and destination pinned in an area with no road network",
     "1. Pick the two pins.\n2. Observe.",
     "OSRM returns no route",
     "Prompt 'Adjust map pins' shown; Create Ride disabled until pins moved.",
     "Medium"],
    ["TC-OR-04", "Departure time less than 5 minutes from now",
     "Driver role active",
     "1. Pick today and time = now + 2 min.\n2. Tap Create Ride.",
     "Time gap = 2 minutes",
     "Inline error 'Departure must be at least 5 minutes in the future.'; Create Ride disabled.",
     "High"],
    ["TC-OR-05", "Available seats above maximum",
     "Driver role active",
     "1. Try to select 5 seats in the segmented selector.",
     "Seats = 5",
     "Selector caps at 4; value 5 cannot be picked through the UI.",
     "Medium"],
    ["TC-OR-06", "No payment methods selected",
     "Driver role active; all other fields valid",
     "1. Uncheck every payment method chip.\n2. Tap Create Ride.",
     "acceptedPaymentMethods = []",
     "Inline error 'Select at least one payment method.'; Create Ride disabled.",
     "High"],
    ["TC-OR-07", "Recurring ride with no day selected",
     "Driver role active",
     "1. Toggle Ride Type to Recurring.\n2. Leave the day chip group empty.\n3. Tap Create Ride.",
     "rideType='recurring'\nrecurringDays=[]",
     "Inline error 'Select at least one day for the recurring ride.'.",
     "High"],
    ["TC-OR-08", "Network drops during ride creation",
     "Driver role active; all fields valid",
     "1. Turn off Wi-Fi and mobile data.\n2. Tap Create Ride.",
     "Network offline",
     "Write is cached by Firestore offline persistence and flushed automatically on reconnect.",
     "Medium"],
    ["TC-OR-09", "Departure date more than 90 days ahead",
     "Driver role active",
     "1. Open the date picker.\n2. Try to scroll to today + 95.",
     "Date offset = +95 days",
     "Date picker disables dates beyond today+90; cannot be selected through the UI.",
     "Low"],
    ["TC-OR-10", "No vehicle registered on profile",
     "Driver role active; users/{uid}/vehicles is empty",
     "1. Open the Offer a Ride screen.",
     "vehicles = []",
     "User is redirected to the Add Vehicle Details screen first; Offer a Ride is blocked.",
     "High"],
    ["TC-OR-11", "User is in Passenger role only",
     "users/{uid}.role = 'passenger'",
     "1. Open the Offer a Ride screen.",
     "role = passenger",
     "Prompt 'Switch to Driver role to offer a ride'; Offer a Ride screen disabled.",
     "Medium"],
    ["TC-OR-12", "Origin equals destination",
     "Driver role active",
     "1. Pick the same lat,lng for both pins.\n2. Tap Create Ride.",
     "Distance = 0 km",
     "Inline error 'Origin and destination must be different.'.",
     "Medium"],
]

OFFER_EP = [
    ["Available Seats (int)", "1, 2, 3, 4", "<= 0; >= 5; non-integer"],
    ["Departure date offset (days from today)", "0 to 90", "< 0 (past); > 90"],
    ["Departure time gap (minutes from now)", ">= 5", "< 5"],
    ["Origin / Destination country", "Pakistan (PK)", "Any non-PK country"],
    ["Distance (km)", "> 0", "= 0 (origin == destination)"],
    ["Payment methods count", "1 to 4", "0"],
    ["Recurring days count (when rideType=recurring)", "1 to 7", "0"],
]

OFFER_BVA = [
    ["Available Seats", "0", "1", "2", "4", "5"],
    ["Departure date offset (days)", "-1", "0 (today)", "30", "90", "91"],
    ["Departure time gap (minutes)", "4", "5", "30", "(bounded by date)", "(none distinct)"],
    ["Distance (km)", "0", "0.1", "12", "(no formal upper)", "(none enforced)"],
    ["Payment methods count", "0", "1", "2", "4", "5 (cap)"],
    ["Recurring days count", "0", "1", "3", "7", "8 (cap)"],
]

RATING_DESC = {
    "id": "UC-4.2",
    "name": "Submit Rating & Review",
    "actors": "Passenger / Driver (Primary), Firebase Firestore (Secondary), Sentiment Classifier (Secondary)",
    "preconditions": [
        "Both users completed a shared ride and the rides document is in status = completed.",
        "The rater (current user) has not previously rated this ride.",
        "The user is authenticated and online.",
    ],
    "postconditions": [
        "A new ratings document is written with stars, optional comment, sentiment tag and createdAt timestamp.",
        "The ratee user's aggregate rating in users/{uid}.rating is recalculated.",
        "The ratee user's trust score is updated based on the sentiment classifier output.",
    ],
    "normal_flow": [
        "System prompts the user with the 'Rate your ride' screen on ride completion.",
        "User taps a star count from 1 to 5.",
        "User optionally types a comment (<= 250 chars).",
        "User taps Submit.",
        "System writes the ratings document to Firestore.",
        "Sentiment classifier tags the comment as positive / neutral / negative and writes it back.",
        "System recalculates users/{rateeId}.rating as the average of all star scores received.",
        "System shows a Thank You confirmation and dismisses the prompt.",
    ],
    "alt": [
        ("1a User skips rating", [
            "In step 1, the user taps 'Skip for now'.",
            "System places a persistent notification on the user's dashboard reminding them to rate later.",
            "Use case terminates.",
        ]),
        ("3a User submits without a comment", [
            "In step 3, the user leaves the comment field empty.",
            "Sentiment is recorded as neutral by default.",
            "Use case resumes at step 4.",
        ]),
    ],
    "exceptions": [
        ("5a Database write failure", [
            "In step 5, if the Firestore write fails, the system shows 'Submission failed. Try again.'.",
            "Use case resumes at step 4.",
        ]),
        ("3b Comment exceeds 250 characters", [
            "If the user pastes a comment longer than 250 characters, the input is hard-truncated at 250 and a helper text 'Max 250 characters' is shown.",
        ]),
        ("Pre-1a Ride not yet completed", [
            "If the ride is still in status active, in_progress or full, the rate prompt is not reachable.",
            "Use case terminates.",
        ]),
    ],
    "quality": "Ratings must be strictly tied to a completed ride so review manipulation is prevented. The sentiment classifier should return within 500 ms. The aggregate rating recalculation must be atomic (transactional) so concurrent ratings cannot corrupt the average.",
}

RATING_TCS = [
    ["TC-RR-01", "Submit a valid 5-star rating with a positive comment",
     "Ride status=completed; rater has not rated this ride",
     "1. Open rate prompt after ride completion.\n2. Tap 5 stars.\n3. Type a short comment.\n4. Tap Submit.",
     "stars=5\ncomment='Great driver, on time.'",
     "ratings doc written; sentiment=positive; ratee aggregate recalculated; Thank You shown.",
     "High"],
    ["TC-RR-02", "Submit a 1-star rating with a negative comment",
     "Ride status=completed",
     "1. Tap 1 star.\n2. Type a negative comment.\n3. Tap Submit.",
     "stars=1\ncomment='Rude behaviour, late by 25 min.'",
     "Saved with sentiment=negative; trust score of ratee decreases.",
     "High"],
    ["TC-RR-03", "Skip rating from prompt",
     "Ride status=completed; rate prompt shown",
     "1. Tap 'Skip for now'.",
     "n/a",
     "Prompt dismissed; persistent reminder placed on dashboard; ratings doc NOT written.",
     "Medium"],
    ["TC-RR-04", "Submit a rating without any comment",
     "Ride status=completed",
     "1. Tap 4 stars.\n2. Leave comment empty.\n3. Tap Submit.",
     "stars=4\ncomment=''",
     "Saved with sentiment=neutral; aggregate recalculated.",
     "Medium"],
    ["TC-RR-05", "Comment exceeds 250 characters",
     "Ride status=completed",
     "1. Paste a 260-character comment.\n2. Tap Submit.",
     "comment length=260",
     "Input is hard-truncated at 250 and 'Max 250 characters' helper text is shown.",
     "High"],
    ["TC-RR-06", "Submit attempted without selecting any star",
     "Rate prompt shown",
     "1. Do NOT tap any star.\n2. Tap Submit.",
     "stars=0 (none selected)",
     "Submit button disabled OR inline error 'Please select a star rating.'.",
     "High"],
    ["TC-RR-07", "Rate before ride completion",
     "Ride status=in_progress",
     "1. Try to access the rate prompt via deep link.",
     "ride.status=in_progress",
     "Access blocked; user is informed the rating is available only after completion.",
     "High"],
    ["TC-RR-08", "Duplicate rating for the same ride",
     "Rater has already rated this ride",
     "1. Open the rate prompt for the same ride again.",
     "existing rating from raterId+rideId",
     "Existing rating shown read-only; submit is not allowed.",
     "High"],
    ["TC-RR-09", "Firestore write failure",
     "Ride status=completed; Firestore unreachable",
     "1. Tap Submit while Firestore is unreachable.",
     "Network blocking Firestore",
     "Error 'Submission failed. Try again.'; user can retry; no doc written.",
     "Medium"],
    ["TC-RR-10", "Profanity in comment",
     "Ride status=completed",
     "1. Type a comment containing a profanity term.\n2. Tap Submit.",
     "comment contains banned word",
     "Saved but flagged; reader sees 'Comment hidden - Show anyway'.",
     "Medium"],
    ["TC-RR-11", "Self-rating attempt (raterId == rateeId)",
     "Direct API call",
     "1. Send a rating where raterId equals rateeId.",
     "raterId == rateeId",
     "API rejects with HTTP 400 and error 'Cannot rate yourself.'.",
     "Low"],
    ["TC-RR-12", "Rate a cancelled ride",
     "ride.status = cancelled",
     "1. Try to access the rate prompt for a cancelled ride.",
     "status=cancelled",
     "Prompt is not reachable; no rating possible.",
     "Medium"],
]

RATING_EP = [
    ["Stars (int)", "1, 2, 3, 4, 5", "0; > 5; negative; non-integer"],
    ["Comment length (chars)", "0 to 250", "> 250"],
    ["Ride status (precondition)", "completed", "active; in_progress; full; cancelled"],
    ["raterId vs rateeId", "raterId != rateeId", "raterId == rateeId"],
    ["Rating uniqueness per ride", "First rating from this rater for this ride", "Second or later rating from same rater for same ride"],
]

RATING_BVA = [
    ["Stars", "0", "1", "3", "5", "6"],
    ["Comment length (chars)", "n/a (0 valid)", "0", "100", "250", "251"],
]


# ---------- per-use-case section builders ----------

def write_use_case_section(doc, desc):
    add_heading(doc, f"3.  Selected Use Case  -  {desc['name']}", level=1)
    add_para(doc,
             "The use case below was first authored in Assignment 2. It is restated "
             "here so this document reads stand-alone before its test cases are derived.")
    add_table_caption(doc, f"Table 3.1  -  Detailed use case specification ({desc['id']})")
    rows = [
        ["Use Case ID", desc["id"]],
        ["Use Case Name", desc["name"]],
        ["Actors", desc["actors"]],
        ["Preconditions", "\n".join(f"{i+1}.  {p}" for i, p in enumerate(desc["preconditions"]))],
        ["Postconditions", "\n".join(f"{i+1}.  {p}" for i, p in enumerate(desc["postconditions"]))],
        ["Normal Flow", "\n".join(f"{i+1}.  {s}" for i, s in enumerate(desc["normal_flow"]))],
        ["Alternative Flows",
         "\n\n".join(f"{title}:\n" + "\n".join(f"   {i+1}.  {s}" for i, s in enumerate(steps))
                     for title, steps in desc["alt"]) or "None"],
        ["Exceptions",
         "\n\n".join(f"{title}:\n" + "\n".join(f"   {i+1}.  {s}" for i, s in enumerate(steps))
                     for title, steps in desc["exceptions"]) or "None"],
        ["Quality Requirements", desc["quality"]],
    ]
    build_table(doc, ["Field", "Details"], rows, col_widths_cm=[3.5, 13.0])


def write_test_cases_section(doc, tcs):
    add_heading(doc, "4.  Use-Case-Based Test Cases", level=1)
    add_para(doc,
             "The test cases below are derived directly from the normal flow, the "
             "alternative flows and the exceptions of the selected use case. Each row "
             "exercises one specific path or condition; together they cover every "
             "numbered scenario in Section 3.")
    add_table_caption(doc, f"Table 4.1  -  Use-case-based test cases ({len(tcs)} total)")
    headers = ["TC ID", "Scenario", "Preconditions", "Test Steps", "Test Data", "Expected Result", "Priority"]
    build_table(doc, headers, tcs, header_size=9, body_size=8,
                col_widths_cm=[1.6, 3.0, 2.6, 3.5, 2.6, 3.4, 1.3])


def write_ep_section(doc, ep):
    add_heading(doc, "5.  Equivalence Partitioning", level=1)
    add_para(doc,
             "For field-level testing of the same core process, each input field is "
             "partitioned into valid and invalid equivalence classes. Selecting one "
             "value from each class is sufficient to expose defects of the same family, "
             "which keeps the number of test cases manageable while still covering "
             "every distinct class.")
    add_table_caption(doc, "Table 5.1  -  Equivalence classes for the selected use case")
    build_table(doc, ["Input Field", "Valid Partition(s)", "Invalid Partition(s)"], ep,
                col_widths_cm=[5.0, 6.0, 6.0])


def write_bva_section(doc, bva):
    add_heading(doc, "6.  Boundary Value Analysis", level=1)
    add_para(doc,
             "Defects tend to cluster at the edges of input ranges. The table below "
             "lists boundary points for each numerical or length-bounded field of the "
             "selected use case. For every field a value just below the lower bound, "
             "exactly the lower bound, a nominal value, exactly the upper bound, and "
             "one value just above the upper bound are tested.")
    add_table_caption(doc, "Table 6.1  -  Boundary points per input field")
    build_table(doc,
                ["Field", "Lower - 1 (invalid)", "Lower bound (valid)", "Nominal (valid)", "Upper bound (valid)", "Upper + 1 (invalid)"],
                bva,
                col_widths_cm=[4.0, 2.8, 2.8, 2.6, 2.8, 2.8])


def write_conclusion(doc, member):
    add_heading(doc, "7.  Conclusion", level=1)
    add_para(
        doc,
        f"This document presents the personal contribution of {member['name']} ({member['reg']}) "
        f"to Group {GROUP_NO}'s Assignment 4 for {COURSE}. The selected core process "
        f"({member['use_case_id']}  -  {member['use_case_name']}) was tested using the "
        "use-case-based testing technique, with explicit coverage of the normal flow, "
        "alternative flows and documented exceptions. Equivalence partitioning and "
        "boundary value analysis were then applied to each input field of the same use "
        "case to provide robust field-level coverage. Together the three deliverables "
        "form a complete test design for this slice of the ShareWay system.",
    )


# ---------- master build ----------

CONTENT_BY_REG = {
    "FA23-BAI-050": (SIGNUP_DESC, SIGNUP_TCS, SIGNUP_EP, SIGNUP_BVA),
    "FA23-BAI-006": (OFFER_DESC, OFFER_TCS, OFFER_EP, OFFER_BVA),
    "FA23-BAI-028": (RATING_DESC, RATING_TCS, RATING_EP, RATING_BVA),
}


def build_for_member(member):
    desc, tcs, ep, bva = CONTENT_BY_REG[member["reg"]]
    doc = Document()
    configure_default_styles(doc)

    cover_page(doc, member)

    section_project_overview(doc)
    section_group_split(doc, member["reg"])
    doc.add_page_break()

    write_use_case_section(doc, desc)
    doc.add_page_break()

    write_test_cases_section(doc, tcs)
    doc.add_page_break()

    write_ep_section(doc, ep)
    write_bva_section(doc, bva)
    write_conclusion(doc, member)

    out_path = os.path.join(OUT_DIR, f"{member['reg']}_Assign4_G{GROUP_NO}.docx")
    doc.save(out_path)
    return out_path


def main():
    for m in MEMBERS:
        print(build_for_member(m))


if __name__ == "__main__":
    main()
